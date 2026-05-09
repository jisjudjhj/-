from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"F:\IDEAwenjian\大数据电商系统")
TEMPLATE_PATH = ROOT / "4-AI工具使用说明（选用模板）（2026年版）.docx"
OUTPUT_PATH = ROOT / "4-AI工具使用说明（近30天-国产AI-已填写）.docx"


ROWS = [
    {
        "seq": "1",
        "tool": "DeepSeek-R1，网页端访问，2026年03月30日 20:10-22:00",
        "purpose": "用户分群建模：协助生成 KMeans 训练主流程、冷启动用户分流、轮廓系数评估与分群摘要字段。",
        "prompt": "请基于电商用户近90天订单与近30天行为特征，生成一个 Python KMeans 训练函数，要求区分冷启动用户，输出 segment_code、silhouette_score、top_categories 和运营建议。",
        "reply": "给出标准化、K 值解析、轮廓系数计算、簇中心回写和冷启动用户单独建段的实现骨架。",
        "manual": "人工改成项目真实字段名，补充 top_categories/top_tags 聚合、confidence_score、segment_name 与 persona_summary 的结果结构，并处理空样本边界。",
        "adoption": "AI 生成内容约占该模块初稿的 40%，经人工重构后最终采纳约 45%。",
        "snippet_path": ROOT / "backend" / "python_analytics" / "kmeans_trainer.py",
        "ranges": [(37, 60), (152, 178)],
        "appendix_title": "序号1的佐证材料：用户分群训练脚本（KMeans + 冷启动）",
        "appendix_note": "注：2026年03月30日使用 DeepSeek-R1 协助输出聚类训练初版，重点用于补齐 K 值选择、冷启动用户分流和分群输出字段；最终实现已结合本项目真实特征列、标签聚合和结果结构做人工重写。",
    },
    {
        "seq": "2",
        "tool": "GLM-4-Plus，网页端访问，2026年04月02日 19:40-21:20",
        "purpose": "分群可视化说明：协助把真实 KMeans 逻辑拆成可答辩的页面结构，突出原理说明、任务管理和结果追踪。",
        "prompt": "请为 Vue 管理端设计一个用户分群页面，要求首屏解释 KMeans 原理，下方展示任务状态、轮廓系数、手动触发与分群规则，文案适合比赛答辩。",
        "reply": "给出原理区、参数事实区、任务历史区和规则说明区的页面分层，以及偏答辩场景的说明文案。",
        "manual": "人工绑定真实 taskRuntime、segmentDecisionRows、轮廓系数和任务历史字段，并删除空泛文案，统一为本项目实际流程表述。",
        "adoption": "AI 生成内容约占该页面结构初稿的 35%，经人工改写后最终采纳约 40%。",
        "snippet_path": ROOT / "management-pc" / "src" / "views" / "admin" / "UserClusterAnalysis.vue",
        "ranges": [(39, 48), (152, 160), (199, 204)],
        "appendix_title": "序号2的佐证材料：用户分群管理端说明页",
        "appendix_note": "注：2026年04月02日使用 GLM-4-Plus 辅助梳理分群页面的信息层级，主要帮助确定“原理解释-任务管理-指标佐证”的展示顺序；最终字段、按钮行为和指标口径均由人工按项目接口改接。",
    },
    {
        "seq": "3",
        "tool": "Qwen2.5-Coder-32B，网页端访问，2026年04月08日 20:00-22:10",
        "purpose": "推荐链路增强：协助生成推荐结果二次重排逻辑，让返回结果更贴近用户偏好品类，同时保留一定探索多样性。",
        "prompt": "请为 Spring Boot 项目写一个 ResponseBodyAdvice，对推荐接口返回的商品列表按用户偏好品类做二次重排；要求识别嵌套列表、保留 Top 品类覆盖并附带重排原因。",
        "reply": "给出 Advice 入口、商品列表识别、品类加权、Top2 偏好保底和原因拼接的实现思路。",
        "manual": "人工补充中文品类关键词、嵌套 Map/List 遍历、分值常量、reason 去重和多列表兼容，避免误伤非推荐接口。",
        "adoption": "AI 生成内容约占该模块骨架的 50%，经人工调参与兼容处理后最终采纳约 55%。",
        "snippet_path": ROOT / "backend" / "src" / "main" / "java" / "com" / "ecommerce" / "config" / "RecommendationResponseAdvice.java",
        "ranges": [(137, 159), (163, 199), (202, 214)],
        "appendix_title": "序号3的佐证材料：推荐结果二次重排逻辑",
        "appendix_note": "注：2026年04月08日使用 Qwen2.5-Coder-32B 协助搭建推荐重排骨架，主要用于生成 Advice 框架和偏好覆盖策略；最终品类词表、分值和嵌套结构兼容均由人工按项目数据格式细化。",
    },
    {
        "seq": "4",
        "tool": "豆包1.5 Pro，网页端访问，2026年04月10日 19:50-21:30",
        "purpose": "实时监控可视化：协助搭建实时热榜、链路状态、积压告警和用户画像联动的运营监控页。",
        "prompt": "请为电商管理端生成一个实时经营监控页面，要求同时展示热榜、Kafka 消费积压、死信告警、链路步骤状态以及单用户实时画像。",
        "reply": "给出总览页与画像页双层结构、链路状态卡片、告警面板和用户画像区的布局建议。",
        "manual": "人工接入实际 overview、monitorData、pipelineSteps、selectedSnapshot 等字段，并调整成比赛展示用的监控总览 + 画像透视结构。",
        "adoption": "AI 生成内容约占页面结构初稿的 45%，经人工接接口和重写文案后最终采纳约 50%。",
        "snippet_path": ROOT / "management-pc" / "src" / "views" / "admin" / "RealtimeStreamBoard.vue",
        "ranges": [(63, 118), (181, 200), (212, 227)],
        "appendix_title": "序号4的佐证材料：实时经营监控页",
        "appendix_note": "注：2026年04月10日使用豆包1.5 Pro 辅助生成实时监控页面结构，重点用于梳理热榜、链路状态和告警面板的呈现方式；最终接口字段、状态判定和实时画像切页逻辑由人工补全。",
    },
    {
        "seq": "5",
        "tool": "Qwen2.5-Coder-32B，网页端访问，2026年04月12日 20:20-22:20",
        "purpose": "推荐预览页：协助整理用户画像、实时分群、策略权重和结果解释的一体化运营视图。",
        "prompt": "请为推荐系统做一个管理端推荐预览页，要求包含自动刷新、用户画像、当前实验组权重、推荐结果解释和页面级标签导航。",
        "reply": "给出查询区、实时刷新提示、页面级 Tabs、实验组说明卡和画像标签区的前端结构建议。",
        "manual": "人工绑定 autoRefresh、PageSectionTabs、realtimeSegment、previewHeroTags 等真实字段，并补上引导流程与页面不可见自动暂停逻辑。",
        "adoption": "AI 生成内容约占该页面首版结构的 55%，经人工细化后最终采纳约 60%。",
        "snippet_path": ROOT / "management-pc" / "src" / "views" / "admin" / "RecommendPreview.vue",
        "ranges": [(87, 96), (108, 115), (210, 220), (612, 612)],
        "appendix_title": "序号5的佐证材料：推荐策略预览页",
        "appendix_note": "注：2026年04月12日使用 Qwen2.5-Coder-32B 协助整理推荐预览页首版结构，主要用于串联“刷新控制-画像洞察-实验权重-结果解释”；最终页面字段映射、引导步骤和刷新策略由人工实现。",
    },
    {
        "seq": "6",
        "tool": "豆包1.5 Pro，网页端访问，2026年04月13日 19:30-21:10",
        "purpose": "A/B 实验分析：协助生成参数对照、相对基线提升和分层差异分析页面，支持答辩时解释实验有效性。",
        "prompt": "请生成一个 Vue A/B 实验分析页，要求展示实验组权重方案、CTR/转化率对照、相对基线提升，以及新客老客和高低客单分层差异。",
        "reply": "给出实验头部指标卡、分组权重说明、相对提升卡片和分层表格的页面组织方式。",
        "manual": "人工补充 groupStrategyMap、overviewLiftRows、分层口径说明和真实指标格式化逻辑，避免只剩通用图表模板。",
        "adoption": "AI 生成内容约占该页面初稿的 42%，经人工绑定真实指标和分层规则后最终采纳约 48%。",
        "snippet_path": ROOT / "management-pc" / "src" / "views" / "admin" / "ABTestPanel.vue",
        "ranges": [(121, 173), (221, 240), (359, 374), (511, 522)],
        "appendix_title": "序号6的佐证材料：A/B 实验分析页",
        "appendix_note": "注：2026年04月13日使用豆包1.5 Pro 辅助生成 A/B 实验页面骨架，重点帮助梳理参数方案、提升幅度和分层对照；最终实验组映射、基线比较和格式化计算均为人工按项目数据重写。",
    },
    {
        "seq": "7",
        "tool": "DeepSeek-R1，网页端访问，2026年04月16日 20:00-22:30",
        "purpose": "经营分析批处理：协助生成 Python 日统计脚本，统一汇总行为漏斗、热力图和销售日表。",
        "prompt": "请用 SQLAlchemy 为电商项目写一个日批统计脚本，要求按日期回写行为日报、漏斗日报、热力图和销售日报，并补齐零值日期、7日均线与模型版本字段。",
        "reply": "给出 run 主流程、事务包裹、行为统计 SQL、销售日表 upsert 和移动平均字段的实现框架。",
        "manual": "人工补充 job_log、模型版本、MySQL ON DUPLICATE KEY UPDATE、日期补齐细节和本项目表结构字段名。",
        "adoption": "AI 生成内容约占该脚本首版的 52%，经人工联调和表结构适配后最终采纳约 58%。",
        "snippet_path": ROOT / "backend" / "python_analytics" / "daily_analytics_main.py",
        "ranges": [(19, 34), (50, 71), (219, 267)],
        "appendix_title": "序号7的佐证材料：经营分析日批脚本",
        "appendix_note": "注：2026年04月16日使用 DeepSeek-R1 辅助生成日批统计脚本的初版，重点补齐多张分析表的 upsert 思路和移动平均计算；最终事务日志、字段名和 MySQL 兼容写法由人工完善。",
    },
    {
        "seq": "8",
        "tool": "GLM-4-Plus，网页端访问，2026年04月17日 19:20-21:00",
        "purpose": "小程序商品详情页：协助整理商品详情、推荐相似商品、秒杀状态和 AI 摘要的多状态逻辑。",
        "prompt": "请为微信小程序商品详情页重构页面状态，要求同时处理商品详情、秒杀倒计时、相似推荐、AI 摘要、评价列表和页面失活保护。",
        "reply": "给出 decisionBoard 状态对象、详情加载主流程、页面失活判断和 AI 摘要触发顺序的代码建议。",
        "manual": "人工改写为小程序 Page 结构，补充 safeSetData、秒杀校验、真实接口路径和评价加载流程，避免在页面离开后继续 setData。",
        "adoption": "AI 生成内容约占该页面重构思路的 38%，经人工改成项目真实小程序逻辑后最终采纳约 42%。",
        "snippet_path": ROOT / "user-miniprogram---2" / "pages" / "product-detail" / "index.js",
        "ranges": [(13, 18), (118, 167)],
        "appendix_title": "序号8的佐证材料：小程序商品详情页（推荐 + 秒杀 + AI 摘要）",
        "appendix_note": "注：2026年04月17日使用 GLM-4-Plus 协助整理商品详情页的多状态处理顺序，主要帮助收敛 decisionBoard、秒杀和 AI 摘要的装配流程；最终代码已人工改成微信小程序可运行写法。",
    },
    {
        "seq": "9",
        "tool": "Qwen2.5-Coder-32B，网页端访问，2026年04月18日 14:10-18:20",
        "purpose": "演示数据生成：协助批量生成比赛演示所需的账号、商家、用户和基础资源 SQL 种子数据。",
        "prompt": "请为电商系统生成可直接导入 MySQL 的 seed.sql，要求拆分基础账号、商家、用户、分类和轮播图数据，字段与 BCrypt 密码格式兼容，并注明自动生成时间。",
        "reply": "给出总种子文件头、基础块说明、统一时间戳、默认密码说明和批量 INSERT 的生成方式。",
        "manual": "人工检查字段顺序、命名风格、类目分布、账号规模和 SQL 分块结构，剔除与本项目表结构不匹配的字段。",
        "adoption": "AI 生成内容约占数据脚本初稿的 30%，经人工校正后最终采纳约 35%。",
        "snippet_path": ROOT / "backend" / "src" / "main" / "resources" / "sql" / "seed.sql",
        "ranges": [(1, 18)],
        "appendix_title": "序号9的佐证材料：比赛演示数据 seed.sql",
        "appendix_note": "注：2026年04月18日使用 Qwen2.5-Coder-32B 协助生成演示数据脚本初稿，重点用于批量账号和基础数据模板；最终表字段、账号命名、类目分布和导入顺序均由人工逐项核对后保留。",
    },
]


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def clear_table_row(row) -> None:
    for cell in row.cells:
        cell.text = ""


def format_cell(value: str) -> str:
    return value.strip()


def extract_snippet(path: Path, ranges: list[tuple[int, int]]) -> str:
    lines = path.read_text(encoding="utf-8").splitlines()
    parts: list[str] = []
    for index, (start, end) in enumerate(ranges):
        if index:
            parts.append("...")
        for line_no in range(start, end + 1):
            content = lines[line_no - 1] if 0 <= line_no - 1 < len(lines) else ""
            parts.append(f"{line_no:>4} | {content}")
    return "\n".join(parts)


def apply_monospace(paragraph: Paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)


def fill_table(doc: Document) -> None:
    table = doc.tables[0]
    for index, row_data in enumerate(ROWS, start=1):
        row = table.rows[index]
        row.cells[0].text = row_data["seq"]
        row.cells[1].text = format_cell(row_data["tool"])
        row.cells[2].text = format_cell(row_data["purpose"])
        row.cells[3].text = format_cell(row_data["prompt"])
        row.cells[4].text = format_cell(row_data["reply"])
        row.cells[5].text = format_cell(f"AI：{row_data['reply']}\n人工：{row_data['manual']}")
        row.cells[6].text = format_cell(row_data["adoption"])

    clear_table_row(table.rows[10])
    clear_table_row(table.rows[11])


def fill_basic_info(doc: Document) -> None:
    doc.paragraphs[2].text = "作品编号：待填写        作品名称：基于大数据分析的电商个性化推荐系统"
    doc.paragraphs[3].text = "说明：本说明仅统计 2026年03月20日至2026年04月18日期间重点模块的 AI 辅助使用情况。"
    doc.paragraphs[4].text = "限制说明：仅使用国产 AI（豆包、DeepSeek、GLM、Qwen），附录仅保留文字与代码片段，不附图片、视频。"
    doc.paragraphs[13].text = "附录2：近30天 AI 使用文字佐证（不含图片、视频）"


def fill_appendix(doc: Document) -> None:
    appendix_title_paragraphs = doc.paragraphs[14:23]
    for paragraph, row_data in zip(appendix_title_paragraphs, ROWS):
        paragraph.text = row_data["appendix_title"]
        current = paragraph

        current = insert_paragraph_after(current, f"对应文件：{row_data['snippet_path'].relative_to(ROOT).as_posix()}")
        location_text = "、".join([f"第 {start}-{end} 行" for start, end in row_data["ranges"]])
        current = insert_paragraph_after(current, f"代码位置：{location_text}")
        current = insert_paragraph_after(current, "代码摘录：")
        code_paragraph = insert_paragraph_after(current)
        apply_monospace(code_paragraph, extract_snippet(row_data["snippet_path"], row_data["ranges"]))
        current = code_paragraph
        current = insert_paragraph_after(current, row_data["appendix_note"])
        current = insert_paragraph_after(current, f"使用模型与时间：{row_data['tool']}")
        current = insert_paragraph_after(current, f"对应提示词：{row_data['prompt']}")
        insert_paragraph_after(current, f"人工修改摘要：{row_data['manual']}")


def main() -> None:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"模板不存在：{TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)
    fill_basic_info(doc)
    fill_table(doc)
    fill_appendix(doc)
    doc.save(OUTPUT_PATH)

    print(
        json.dumps(
            {
                "status": "ok",
                "output": str(OUTPUT_PATH),
                "rows": len(ROWS),
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
